#!/usr/bin/env python3
"""Generate one EE Report Summary document per sport from the metrics merge CSV."""

from __future__ import annotations

import re
from pathlib import Path

import pandas as pd
from docx import Document

BASE_DIR = Path(__file__).resolve().parent
TEMPLATE_PATH = BASE_DIR / "EE Report Summary - Template.docx"
CSV_PATH = BASE_DIR / "EE_Sport_Metrics_2026(DATA MERGE).csv"
OUTPUT_DIR = BASE_DIR / "generated_ee_reports"


def normalize_key(value: str) -> str:
    """Normalize a key for resilient matching between template labels and CSV headers."""
    text = (value or "").strip().lower()
    text = (
        text.replace("→", "?")
        .replace("�", "?")
        .replace(">=", "?")
        .replace("≥", "?")
        .replace("–", "-")
        .replace("—", "-")
    )
    text = re.sub(r"\s+", " ", text)
    text = re.sub(r"[^a-z0-9%?()/:._ -]", "", text)
    return text


def format_value(value: str) -> str:
    """Format plain numeric values to one decimal place."""
    text = (value or "").strip()
    if not text:
        return text

    # Keep currency-like values as-is.
    if "$" in text:
        return text

    try:
        number = float(text)
    except ValueError:
        return text

    if abs(number) < 0.05:
        number = 0.0
    return f"{number:.1f}"


def build_row_mapping(row: pd.Series) -> tuple[dict[str, str], dict[str, str]]:
    raw_map: dict[str, str] = {}
    norm_map: dict[str, str] = {}

    for key, value in row.items():
        if pd.isna(value):
            continue
        key_str = str(key).strip()
        value_str = format_value(str(value).strip())
        raw_map[key_str] = value_str
        norm_map[normalize_key(key_str)] = value_str

    aliases = {
        "FY26Profle": "FY26Profile",
        "Rank Profile CategoryRank": "Rank Profile Category",
        "% Athletes Fitness Tested ≥2x/yr": "% Athletes Fitness Tested ?2x/yr",
        "% Athletes Fitness Tested ≥2x/yr70th Perc": "% Athletes Fitness Tested ?2x/yr70th Perc",
        "Total Conversion Prov→Nat (4y)": "Total Conversion Prov?Nat (4y)",
        "% Avg Conversion Prov→Nat (4y)70th Perc": "% Avg Conversion Prov?Nat (4y)70th Perc",
        "Avg Years Targeted (Prov Dev, 2025–26)": "Avg Years Targeted (Prov Dev, 2025�26)",
        "Avg Years Targeted (Prov Dev, 2025–26)70th Perc": "Avg Years Targeted (Prov Dev, 2025�26)70th Perc",
    }

    for template_key, csv_key in aliases.items():
        if csv_key in raw_map:
            raw_map[template_key] = raw_map[csv_key]
            norm_map[normalize_key(template_key)] = raw_map[csv_key]

    return raw_map, norm_map


def replacement_for_text(text: str, sport_name: str, raw_map: dict[str, str], norm_map: dict[str, str]) -> str:
    if not text:
        return text

    replaced = text.replace("<SPORT>", sport_name)
    stripped = replaced.strip()

    if stripped in raw_map:
        return raw_map[stripped]

    norm = normalize_key(stripped)
    if norm in norm_map:
        return norm_map[norm]

    return replaced


def update_document(doc: Document, sport_name: str, raw_map: dict[str, str], norm_map: dict[str, str]) -> None:
    for paragraph in doc.paragraphs:
        new_text = replacement_for_text(paragraph.text, sport_name, raw_map, norm_map)
        if new_text != paragraph.text:
            paragraph.text = new_text

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    new_text = replacement_for_text(paragraph.text, sport_name, raw_map, norm_map)
                    if new_text != paragraph.text:
                        paragraph.text = new_text


def safe_filename(name: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9._ -]+", "", name).strip()
    cleaned = re.sub(r"\s+", " ", cleaned)
    cleaned = cleaned.replace(" ", "_")
    return cleaned or "unknown_sport"


def main() -> None:
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"Template not found: {TEMPLATE_PATH}")
    if not CSV_PATH.exists():
        raise FileNotFoundError(f"CSV not found: {CSV_PATH}")

    df = pd.read_csv(CSV_PATH, header=1, dtype=str, encoding="cp1252")

    # Keep only rows with a sport value in the primary Sport column.
    df = df[df["Sport"].notna() & (df["Sport"].str.strip() != "")].copy()

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    generated = 0
    for _, row in df.iterrows():
        sport_name = str(row["Sport"]).strip()
        raw_map, norm_map = build_row_mapping(row)

        doc = Document(TEMPLATE_PATH)
        update_document(doc, sport_name, raw_map, norm_map)

        out_path = OUTPUT_DIR / f"EE_Report_Summary_{safe_filename(sport_name)}.docx"
        doc.save(out_path)
        generated += 1

    print(f"Generated {generated} report files in: {OUTPUT_DIR}")


if __name__ == "__main__":
    main()
